import os
from PIL import Image
import cv2
from docx import Document
from docx.shared import Inches
import re

BASE = r'D:\_Captura_Imagens_Video\\'

def str_time_to_seconds(t: str):
    t = t.replace(",", ".")
    pats = re.findall(r"\d+", t)
    if len(pats) >= 3:
        h, m, s = int(pats[0]), int(pats[1]), float(pats[2])
        return h*3600 + m*60 + s
    return 0

def parse_transcricao(path):
    blocos = []
    with open(path, encoding="utf-8") as f:
        linhas = [l.strip() for l in f if l.strip()]
    i = 0
    while i < len(linhas):
        if re.match(r'\d{2}:\d{2}:\d{2}', linhas[i]):
            tempo = linhas[i]
            texto = linhas[i+1] if (i+1)<len(linhas) else ""
            ts = str_time_to_seconds(tempo)
            blocos.append((ts, texto))
            i += 2
        else:
            i += 1
    return blocos

def extrair_multi_frames(video_path, ts, save_dir, bloco_idx, tempo_total, antes, depois):
    cap = cv2.VideoCapture(video_path)
    frames = []
    nomes = []
    offsets = list(range(-antes, depois+1))
    for off in offsets:
        target = ts + off
        if target < 0 or target > tempo_total:
            continue
        label = (
            f"antes{abs(off)}s" if off < 0
            else ("central" if off == 0 else f"depois{off}s")
        )
        cap.set(cv2.CAP_PROP_POS_MSEC, target*1000)
        ret, frame = cap.read()
        if ret:
            img_name = f"frame_bloco{bloco_idx:03d}_{label}_{int(target)}s.jpg"
            img_path = os.path.join(save_dir, img_name)
            Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)).save(img_path)
            frames.append(img_path)
            nomes.append(label)
    cap.release()
    return frames, nomes

def main():
    nome_base = input("Nome base dos arquivos (sem extensão, ex: Tutorial-Cities_Skyline-01): ").strip()
    ext_video = input("Extensão do vídeo (ex: mp4): ").strip().lower()
    ext_txt = input("Extensão da transcrição (ex: txt): ").strip().lower()
    segundos_antes = int(input("Quantos segundos ANTES capturar? (ex: 5): ").strip())
    segundos_depois = int(input("Quantos segundos DEPOIS capturar? (ex: 5): ").strip())

    VIDEO = os.path.join(BASE, nome_base + "." + ext_video)
    TXT = os.path.join(BASE, nome_base + "." + ext_txt)
    pasta_imgs = os.path.join(BASE, f"Imagens_{nome_base}")
    os.makedirs(pasta_imgs, exist_ok=True)

    cap = cv2.VideoCapture(VIDEO)
    tempo_total = cap.get(cv2.CAP_PROP_FRAME_COUNT) / cap.get(cv2.CAP_PROP_FPS)
    cap.release()

    blocos = parse_transcricao(TXT)

    doc = Document()
    doc.add_heading(f"Tutorial: {nome_base}", level=1)
    doc.add_paragraph(f"Arquivo de vídeo: {nome_base}.{ext_video}")
    doc.add_paragraph(f"Arquivo de transcrição: {nome_base}.{ext_txt}")

    for idx, (ts, texto) in enumerate(blocos):
        doc.add_heading(f"Bloco {idx+1} - {ts:.1f}s", level=2)
        doc.add_paragraph(texto)
        frames, labels = extrair_multi_frames(
            VIDEO, ts, pasta_imgs, idx+1, tempo_total, segundos_antes, segundos_depois
        )
        for frame, label in zip(frames, labels):
            doc.add_paragraph(f"Frame {label}:")
            doc.add_picture(frame, width=Inches(3.2))

    # Salva o Word dentro da pasta de imagens!
    word_name = os.path.join(pasta_imgs, f'Resumo_{nome_base}.docx')
    doc.save(word_name)
    print(f"Word criado em: {word_name}")

if __name__ == '__main__':
    main()
